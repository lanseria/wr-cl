"""Test configuration and fixtures for pytest."""
import os
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
# 测试配置和固件


@pytest.fixture
def test_config():
    """Return a test configuration dictionary."""
    return {
        "replacements": {
            "pattern_type": "plain",
            "rules": [
                {
                    "old_text": "公司A",
                    "new_text": "DeepSeek",
                    "options": {
                        "case_sensitive": True,
                        "whole_word": True,
                        "preserve_format": True
                    }
                }
            ]
        },
        "file_settings": {
            "input_path": "./tests/test_docs",
            "file_types": [".docx"],
            "output_path": "./tests/output"
        },
        "advanced": {
            "max_workers": 1,
            "timeout": 30
        }
    }


@pytest.fixture
def sample_docx():
    """Create a sample Word document for testing."""
    doc_path = Path("./tests/test_docs/test.docx")
    doc_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Add a simple paragraph
    p1 = doc.add_paragraph("这是一个关于公司A的测试文档。")
    run = p1.runs[0]
    run.font.size = Pt(12)
    run.font.bold = True

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "公司A"
    table.cell(0, 1).text = "描述"
    table.cell(1, 0).text = "位置"
    table.cell(1, 1).text = "公司A总部"

    doc.save(doc_path)
    return doc_path


@pytest.fixture
def output_dir():
    """Create and return the output directory path."""
    output_path = Path("./tests/output")
    output_path.mkdir(parents=True, exist_ok=True)
    return output_path
