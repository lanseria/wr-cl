"""Test cases for document processor."""
import pytest
from pathlib import Path
from docx import Document
from src.processor import DocumentProcessor
# 文档处理测试


def test_document_processor_initialization(test_config):
    """Test processor initialization."""
    processor = DocumentProcessor(test_config, dry_run=False)
    assert processor.pattern_type == "plain"
    assert len(processor.rules) == 1
    assert processor.rules[0]["old_text"] == "公司A"


def test_document_processing(test_config, sample_docx, output_dir):
    """Test document processing with actual file."""
    processor = DocumentProcessor(test_config, dry_run=False)
    processor.process_document(sample_docx, output_dir)

    # Check the output file
    output_file = output_dir / sample_docx.name
    assert output_file.exists()

    # Verify content changes
    doc = Document(output_file)
    assert "DeepSeek" in doc.paragraphs[0].text
    assert "公司A" not in doc.paragraphs[0].text

    # Check table content
    table = doc.tables[0]
    assert table.cell(0, 0).text == "DeepSeek"
    assert table.cell(1, 1).text == "DeepSeek总部"


def test_dry_run_mode(test_config, sample_docx, output_dir):
    """Test dry run mode."""
    processor = DocumentProcessor(test_config, dry_run=True)
    processor.process_document(sample_docx, output_dir)

    # Check that no output file was created
    output_file = output_dir / sample_docx.name
    assert not output_file.exists()


def test_format_preservation(test_config, sample_docx, output_dir):
    """Test that text formatting is preserved."""
    processor = DocumentProcessor(test_config, dry_run=False)
    processor.process_document(sample_docx, output_dir)

    # Check the output file
    output_file = output_dir / sample_docx.name
    doc = Document(output_file)

    # Verify formatting is preserved
    run = doc.paragraphs[0].runs[0]
    assert run.font.bold is True
    assert run.font.size.pt == 12
