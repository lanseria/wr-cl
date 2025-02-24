"""Document processing functionality."""
from src.logger_config import setup_logger
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Dict, Any, List
from docx.text.paragraph import Paragraph
from docx import Document
from docx.enum.section import WD_HEADER_FOOTER


class DocumentProcessor:
    """Handles the processing of Word documents, including replacing text in headers and footers."""

    def __init__(self, config: Dict[str, Any], dry_run: bool = False):
        """Initialize the document processor."""
        self.config = config
        self.dry_run = dry_run
        self.replacements = config["replacements"]
        self.file_settings = config["file_settings"]
        self.advanced = config["advanced"]
        self.pattern_type = self.replacements["pattern_type"]
        self.rules = self.replacements["rules"]

        # Set the log level from config if available, otherwise use default 'debug'
        self.logger = setup_logger(
            "src.processor", level=config.get("log_level", "debug")
        )

    def _get_files_to_process(self, input_path: Path) -> List[Path]:
        """Get list of files to process based on configured file types."""
        files = []
        for file_type in self.file_settings["file_types"]:
            if not file_type.startswith('.'):
                file_type = f'.{file_type}'
            files.extend(input_path.rglob(f'*{file_type}'))

        self.logger.info(
            f"Found {len(files)} files to process in {input_path}")
        for file in files:
            self.logger.debug(f"Found file: {file}")
        return files

    def process_all(self) -> None:
        """Process all documents in the input directory."""
        input_path = Path(self.file_settings["input_path"])
        output_path = Path(self.file_settings["output_path"])

        if not input_path.exists():
            raise FileNotFoundError(f"Input path does not exist: {input_path}")

        if not self.dry_run:
            output_path.mkdir(parents=True, exist_ok=True)

        files_to_process = self._get_files_to_process(input_path)
        if not files_to_process:
            self.logger.warning("No files found to process")
            return

        with ThreadPoolExecutor(max_workers=self.advanced["max_workers"]) as executor:
            future_to_file = {
                executor.submit(self.process_document, file_path, output_path): file_path
                for file_path in files_to_process
            }

            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    future.result(timeout=self.advanced["timeout"])
                except Exception as e:
                    self.logger.error(
                        f"Error processing {file_path}: {str(e)}")

    def process_document(self, file_path: Path, output_path: Path) -> None:
        """Process a single document, including headers and footers."""
        self.logger.info(f"Processing document: {file_path}")

        if self.dry_run:
            self.logger.info(f"Dry run - would process {file_path}")
            self._preview_changes(file_path)
            return

        try:
            doc = Document(str(file_path))
            modified = False

            # Process main document paragraphs
            for paragraph in doc.paragraphs:
                if self._process_paragraph(paragraph):
                    modified = True

            # Process tables in the document body
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if self._process_paragraph(paragraph):
                                modified = True

            # 处理所有分节（文献[9]）
                for i, section in enumerate(doc.sections):
                    if i > 0 and not section.different_first_page_header_footer:
                        section.header.is_linked_to_previous = False
                        section.footer.is_linked_to_previous = False

                    if self._process_header_footer(section):
                        modified = True

            # Save the document regardless; adjust if you wish to save only modified files.
            output_file = output_path / file_path.name
            # 处理完成后需强制刷新页面布局：
            # doc.element.xml = doc.element.xml.replace(
            #     b'<w:view>normal</w:view>',
            #     b'<w:view>print</w:view>'
            # )
            doc.save(str(output_file))
            if modified:
                self.logger.info(f"Saved modified document to {output_file}")
            else:
                self.logger.info(f"No changes needed for {file_path}")

        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            raise

    def _process_paragraph(self, paragraph: Paragraph) -> bool:
        """带格式保留的文本替换"""
        original_text = paragraph.text
        if not original_text:
            return False

        # 记录所有run的格式和位置
        runs_data = []
        for run in paragraph.runs:
            runs_data.append({
                "text": run.text,
                "font": {
                    "name": run.font.name,
                    "size": run.font.size,
                    "bold": run.font.bold,
                    "italic": run.font.italic,
                    "color": run.font.color.rgb if run.font.color else None
                }
            })

        # 合并所有文本进行替换
        full_text = ''.join([rd["text"] for rd in runs_data])
        modified_text = full_text
        for rule in self.rules:
            modified_text = modified_text.replace(
                rule["old_text"], rule["new_text"])

        if modified_text == full_text:
            return False

        # 重建带格式的runs（文献[5]）
        paragraph.clear()
        new_run = paragraph.add_run(modified_text)

        # 继承第一个run的格式（可根据需要调整）
        if runs_data:
            original_font = runs_data[0]["font"]
            new_run.font.name = original_font["name"]
            new_run.font.size = original_font["size"]
            new_run.font.bold = original_font["bold"]
            new_run.font.italic = original_font["italic"]
            if original_font["color"]:
                new_run.font.color.rgb = original_font["color"]

        return True

    def _process_header_footer(self, section) -> bool:
        """处理所有类型页眉页脚"""
        modified = False

        # 处理三种页眉类型（文献[8]）
        for header_type in [WD_HEADER_FOOTER.PRIMARY,
                            WD_HEADER_FOOTER.FIRST_PAGE,
                            WD_HEADER_FOOTER.EVEN_PAGE]:
            header = section.header if header_type == WD_HEADER_FOOTER.PRIMARY \
                else section.first_page_header if header_type == WD_HEADER_FOOTER.FIRST_PAGE \
                else section.even_page_header

            if header:
                # 处理表格中的页眉内容（文献[4]）
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if self._process_paragraph(paragraph):
                                    modified = True

                # 处理普通段落
                for paragraph in header.paragraphs:
                    if self._process_paragraph(paragraph):
                        modified = True

        # 相同逻辑处理页脚
        for footer_type in [WD_HEADER_FOOTER.PRIMARY,
                            WD_HEADER_FOOTER.FIRST_PAGE,
                            WD_HEADER_FOOTER.EVEN_PAGE]:
            footer = section.footer if footer_type == WD_HEADER_FOOTER.PRIMARY \
                else section.first_page_footer if footer_type == WD_HEADER_FOOTER.FIRST_PAGE \
                else section.even_page_footer

            if footer:
                # 处理表格型页脚
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if self._process_paragraph(paragraph):
                                    modified = True

                for paragraph in footer.paragraphs:
                    if self._process_paragraph(paragraph):
                        modified = True

        return modified

    def _preview_changes(self, file_path: Path) -> None:
        """Preview changes that would be made to the document."""
        try:
            doc = Document(str(file_path))
            changes = []

            # Check main document paragraphs
            for paragraph in doc.paragraphs:
                if not paragraph.text:
                    continue
                original_text = paragraph.text
                for rule in self.rules:
                    old_text = rule["old_text"]
                    new_text = rule["new_text"]
                    if old_text in original_text:
                        changes.append({
                            "location": "body",
                            "old_text": old_text,
                            "new_text": new_text,
                            "context": original_text
                        })

            # Check headers and footers
            for section in doc.sections:
                for part, name in [(section.header, "header"), (section.footer, "footer")]:
                    if part:
                        for paragraph in part.paragraphs:
                            if not paragraph.text:
                                continue
                            original_text = paragraph.text
                            for rule in self.rules:
                                old_text = rule["old_text"]
                                new_text = rule["new_text"]
                                if old_text in original_text:
                                    changes.append({
                                        "location": name,
                                        "old_text": old_text,
                                        "new_text": new_text,
                                        "context": original_text
                                    })

            if changes:
                self.logger.info(f"\nPreview of changes for {file_path}:")
                for i, change in enumerate(changes, 1):
                    self.logger.info(f"\nChange {i} in {change['location']}:")
                    self.logger.info(f"Old text: {change['old_text']}")
                    self.logger.info(f"New text: {change['new_text']}")
                    self.logger.info(f"Context: {change['context']}")
            else:
                self.logger.info(f"No changes would be made to {file_path}")

        except Exception as e:
            self.logger.error(
                f"Error previewing changes for {file_path}: {str(e)}")
            raise
